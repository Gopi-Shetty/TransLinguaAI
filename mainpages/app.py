from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
import os
import PyPDF2
import pdfplumber
from docx import Document
from groq import Groq
from flask_cors import CORS


# Configure the app and allowed file types
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['ALLOWED_EXTENSIONS'] = {'txt', 'pdf', 'docx'}

# Ensure the upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


client = Groq(
    api_key=os.getenv("GROQ_API_KEY")
) #<-- Replace with your actual key

# Enable CORS
CORS(app)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/')
def home():
    return render_template('index.html') #Corrected template name
@app.route('/login')
def login():
    return render_template('log.html')

@app.route('/signup')
def signup():
    return render_template('sign.html')

@app.route('/translator')
def translator():
    return render_template('index1.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        try:
            file.save(filepath)
            extracted_text = extract_text(filepath, filename)
            #Removed AI cleaning as it was causing issues.  Consider re-adding with improved error handling if needed.
            return jsonify({"text": extracted_text})
        except Exception as e:
            return jsonify({"error": f"Error processing upload: {str(e)}"}), 500

    return jsonify({"error": "File type not allowed"}), 400

def extract_text(filepath, filename):
    ext = filename.split('.')[-1].lower()
    try:
        if ext == "txt":
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        elif ext == "pdf":
            return extract_text_from_pdf(filepath)
        elif ext == "docx":
            return extract_text_from_docx(filepath)
        else:
            return "Unsupported file format"
    except Exception as e:
        return f"Error reading file: {str(e)}"

def extract_text_from_pdf(filepath):
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    except Exception:
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
    return text.strip() if text else "No readable text found in PDF"

def extract_text_from_docx(filepath):
    doc = Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def create_docx(text):
    try:
        document = Document()
        document.add_paragraph(text)
        temp_docx_path = "temp.docx"
        document.save(temp_docx_path)
        with open(temp_docx_path, "rb") as f:
            docx_data = f.read()
        os.remove(temp_docx_path)
        return docx_data
    except Exception as e:
        return f"Error creating DOCX: {str(e)}"



@app.route('/translate', methods=['POST'])
def translate_text():

    data = request.json

    text = data.get("text", "")
    target_language = data.get("targetLanguage", "English")
    output_structure = data.get("outputStructure", "formal and educational")

    if not text:
        return jsonify({"error": "Text is required"}), 400

    try:

        prompt = f"""
Translate the following text into {target_language}.

Tone: {output_structure}

Rules:
- If input is a single word, output only one translated word.
- If input is a sentence, preserve its meaning.
- Return ONLY the translated text.

Text:
{text}
"""

        response = client.chat.completions.create(
            model="openai/gpt-oss-120b",
           messages=[
    {
        "role": "system",
        "content": "You are a professional multilingual translator. Translate accurately. Return ONLY the translated text without explanations."
    },
    {
        "role": "user",
        "content": prompt
    }
],
            temperature=0.3,
            max_completion_tokens=2048,
            top_p=1
        )

        translation = response.choices[0].message.content.strip()

        if not translation:
            return jsonify({
                "error": "No translation returned."
            }), 500

        return jsonify({
            "translation": translation
        })

    except Exception as e:
        print(e)

        return jsonify({
            "error": str(e)
        }), 500


if __name__ == "__main__":
    app.run(debug=True)
